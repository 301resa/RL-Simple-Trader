"""
Generate RL Trading Strategy Word document.
Run: /c/ProgramData/anaconda3/python.exe docs/generate_strategy_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
    p.add_run(text)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue header fill
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    return t

# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RL Order Zone Trading Strategy")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Architecture, Entry Logic, Trade Management & Exit Rules").font.size = Pt(13)

doc.add_paragraph()
inst = doc.add_paragraph()
inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
inst.add_run("Instrument: ES / NQ Micro Futures  |  Timeframe: 5-Minute Bars").font.size = Pt(11)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
h1("1. Strategy Overview")
body(
    "This strategy uses a Recurrent PPO (LSTM) reinforcement learning agent to trade "
    "ES/NQ micro futures on 5-minute bars. The agent is trained to enter positions only "
    "at high-confluence Order Zones — price areas where supply/demand imbalance, "
    "liquidity sweeps, and rejection candles align simultaneously."
)
body(
    "The agent operates within a fully defined Markov Decision Process (MDP): one "
    "episode equals one RTH session (~78 bars). At each bar it selects from five "
    "possible actions, subject to hard safety masks."
)

doc.add_paragraph()
h2("Core Design Philosophy")
bullet("Risk is defined BEFORE entry — stop is placed beyond the zone, target is a function of remaining ATR.")
bullet("R-multiples are the unit of measurement — P&L, rewards, and all thresholds are denominated in R.")
bullet("Process quality is rewarded — the agent learns to prefer high-confluence setups, not just profitable ones.")
bullet("Hard guardrails cannot be overridden by the agent — masking enforces non-negotiable risk rules.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 2. ACTION SPACE
# ═════════════════════════════════════════════════════════════════════════════
h1("2. Action Space (5 Actions)")
body("At every 5-minute bar the agent selects exactly one of the following actions:")
doc.add_paragraph()

add_table(
    ["Action", "Code", "Description"],
    [
        ["HOLD / WAIT",    "0", "Do nothing — stay flat or hold current position without change"],
        ["ENTER SHORT",    "1", "Open a short position at current bar close price"],
        ["ENTER LONG",     "2", "Open a long position at current bar close price"],
        ["EXIT",           "3", "Close the current position immediately at market price"],
        ["TRAIL STOP",     "4", "Activate / advance the trailing stop to protect open profits"],
    ]
)

doc.add_paragraph()
body(
    "Actions are masked before the agent samples — invalid actions are removed from "
    "the probability distribution entirely, so the agent never 'tries' a forbidden action."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3. ENTRY CONDITIONS
# ═════════════════════════════════════════════════════════════════════════════
h1("3. Entry Logic")

h2("3.1 Hard Entry Masks (absolute blocks)")
body("These conditions completely prevent entry — they cannot be overridden by the agent:")
doc.add_paragraph()

add_table(
    ["Condition", "Rule"],
    [
        ["Position already open",     "Only one trade open at a time — no pyramiding"],
        ["Max trades per day reached", "Default: 10 trades per session maximum"],
        ["Loss streak pause",          "After 3 consecutive losses, entries blocked for 6 bars"],
        ["End of session",             "No new entries in the last 3 bars of RTH"],
        ["Max drawdown breached",      "If daily loss ≥ 3R, only HOLD is allowed"],
    ]
)

doc.add_paragraph()
h2("3.2 Entry Quality Scoring (via Reward Function)")
body(
    "When entry masks allow an entry, the agent is incentivised to enter at high-quality "
    "setups through bonuses and penalties applied immediately at entry:"
)
doc.add_paragraph()

add_table(
    ["Component", "Value (R)", "Trigger"],
    [
        ["Full Order Zone Confluence",  "+0.50", "All three pillars present: zone + liquidity + rejection"],
        ["In Supply/Demand Zone",       "+0.20", "Price inside a valid S/D zone"],
        ["Liquidity Sweep Present",     "+0.15", "Recent liquidity sweep detected"],
        ["Rejection Candle Present",    "+0.10", "Pin bar or engulfing candle confirmed"],
        ["ATR Has Room",                "+0.15", "ATR not yet in warning/exhaustion territory"],
        ["High R:R Ratio",              "+0.20", "Setup R:R ≥ minimum threshold"],
        ["Entry Cost",                  "−0.005","Transaction cost proxy — applied on every entry"],
        ["No Zone Present",             "−0.05", "Entry attempted with no valid zone"],
        ["ATR Exhausted",               "−0.40", "Session range already used ≥ 95% of daily ATR"],
        ["R:R Below Minimum",           "−0.08", "Risk/reward ratio below 1.5:1"],
        ["No Liquidity Confluence",     "−0.05", "No liquidity sweep or sweep evidence"],
        ["Overtrading",                 "−0.20", "Entry after daily trade limit approached"],
    ]
)

doc.add_paragraph()
h2("3.3 Three Pillars of a Valid Order Zone Entry")
body("An Order Zone requires all three pillars to receive the full +0.50 confluence bonus:")
doc.add_paragraph()
bullet("Supply/Demand Zone", bold_prefix="Pillar 1")
doc.paragraphs[-1].add_run(
    " — A consolidation followed by an impulse move. The zone is the consolidation area. "
    "Zones expire after 200 bars or 3 touches."
)
bullet("Liquidity Sweep", bold_prefix="Pillar 2")
doc.paragraphs[-1].add_run(
    " — Price runs a swing high/low (clearing stop orders) then reverses. "
    "The wick must be ≥ 3% of ATR."
)
bullet("Rejection Candle", bold_prefix="Pillar 3")
doc.paragraphs[-1].add_run(
    " — Either a pin bar (wick ratio ≥ 2×) or an engulfing candle (body ratio ≥ 1.1×) "
    "forming at the zone boundary."
)

doc.add_paragraph()
h2("3.4 Stop Loss Placement")
body("Stop loss is placed beyond the zone boundary, not a fixed number of points:")
doc.add_paragraph()
bullet("Method: Zone-based — stop placed just outside the consolidation area")
bullet("Buffer: 3% of daily ATR beyond the zone top (long) or zone bottom (short)")
bullet("Rule: Stop can NEVER be widened after entry (hard rule in position manager)")
bullet("Example — Long: zone base at 5200, ATR = 100 pts → stop at 5200 − (100 × 0.03) = 5197")

doc.add_paragraph()
h2("3.5 Take Profit Calculation")
body("Take profit is dynamically sized using remaining ATR at the time of entry:")
doc.add_paragraph()
bullet("Method: ATR-remaining — target = 75% of the daily ATR not yet consumed by the session range")
bullet("Minimum R:R: 1.5:1 required (risk/reward must be at least 1.5 before entry is scored favourably)")
bullet("Example: ATR = 100 pts, session range used 40 pts, remaining = 60 pts → target distance = 60 × 0.75 = 45 pts")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 4. TRADE MANAGEMENT
# ═════════════════════════════════════════════════════════════════════════════
h1("4. Trade Management While in Position")

h2("4.1 Position Sizing")
body("Contracts are dynamically sized to risk a fixed percentage of account equity per trade:")
doc.add_paragraph()
add_table(
    ["Parameter", "Value", "Description"],
    [
        ["Risk per trade",  "1% of account",    "Fixed fractional sizing"],
        ["Min contracts",   "1",                 "Floor — always at least 1 contract"],
        ["Max contracts",   "3",                 "Cap — never more than 3 contracts"],
        ["Point value",     "$2.00 (MES/MNQ)",  "Dollar value per point per contract"],
    ]
)
doc.add_paragraph()
body(
    "Formula: n_contracts = floor(account × 0.01 / (stop_distance_pts × point_value)), "
    "clamped to [1, 3]."
)

doc.add_paragraph()
h2("4.2 Trailing Stop Logic")
body(
    "The agent can activate a trailing stop by selecting the TRAIL STOP action. "
    "Once active, the stop advances automatically every bar the action is maintained."
)
doc.add_paragraph()

add_table(
    ["Stage", "Trigger", "Behaviour"],
    [
        ["Activation",   "Unrealised P&L ≥ 2R AND agent selects TRAIL STOP",
         "Trailing stop activates; initial trail distance = ATR × 25%"],
        ["Normal trail", "Each bar agent selects TRAIL STOP",
         "Stop ratchets up (long) / down (short) — never reverses"],
        ["Aggressive",   "Unrealised P&L ≥ 4R",
         "Trail tightens further — agent rewarded for selecting trail here"],
        ["Locked in",    "Trail activates",
         "Minimum 2R is locked in — stop will never fall below entry + 2R"],
    ]
)

doc.add_paragraph()
body(
    "TRAIL STOP is masked (unavailable) when unrealised P&L < 2R — the agent cannot "
    "trail prematurely and risk tightening the stop before the trade has room."
)

doc.add_paragraph()
h2("4.3 Same-Bar Ambiguity Rule")
body(
    "When a single candle's range touches BOTH the stop loss and take profit simultaneously, "
    "the strategy always assumes the STOP was hit first (worst-case fill). "
    "This prevents the agent exploiting ambiguous candles as guaranteed take-profits."
)
bullet("Long: if bar low ≤ stop AND bar high ≥ target → stop loss (loss recorded)")
bullet("Short: if bar high ≥ stop AND bar low ≤ target → stop loss (loss recorded)")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 5. EXIT CONDITIONS
# ═════════════════════════════════════════════════════════════════════════════
h1("5. Exit Conditions (Priority Order)")
body("Exits are processed in strict priority order each bar:")
doc.add_paragraph()

add_table(
    ["Priority", "Exit Reason", "Trigger", "Exit Price"],
    [
        ["1 (highest)", "Stop Loss",     "Bar low ≤ stop (long) or bar high ≥ stop (short)",        "Stop price"],
        ["1 (highest)", "Trailing Stop", "Bar crosses trailing stop level",                          "Trailing stop price"],
        ["2",           "Take Profit",   "Bar high ≥ target (long) or bar low ≤ target (short)",     "Target price"],
        ["3",           "Agent Exit",    "Agent selects EXIT action",                                 "Current bar close"],
        ["4",           "Session End",   "Last bar of RTH session reached with open position",       "Bar close"],
        ["5",           "Max Drawdown",  "Total daily loss (realised + unrealised) ≥ 3R",            "Market close"],
    ]
)

doc.add_paragraph()
h2("Exit Rewards and Penalties")
add_table(
    ["Component", "Value (R)", "Trigger"],
    [
        ["Trailing stop correctly used",   "+0.30", "Position exited via trailing stop"],
        ["Aggressive trail at 4R",         "+0.20", "Trail activated with unrealised ≥ 4R"],
        ["Exit at ATR target",             "+0.20", "Take profit hit at the ATR-based target"],
        ["Held past 4R without trail",     "−0.50", "Position held beyond 4R unrealised without trailing"],
        ["Gave back profit",               "−variable", "Weight × R given back from peak unrealised"],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 6. RISK GUARDRAILS
# ═════════════════════════════════════════════════════════════════════════════
h1("6. Risk Guardrails")

add_table(
    ["Guardrail", "Limit", "Consequence"],
    [
        ["Daily loss limit",      "−3R",          "Trading stops for the session; only HOLD allowed"],
        ["Max trades per day",    "10",            "Entry actions masked for remainder of session"],
        ["Loss streak pause",     "3 consecutive", "Entry blocked for 6 bars; then unblocked"],
        ["No late entries",       "Last 3 bars",   "No new entries in last 15 minutes of RTH"],
        ["ATR exhaustion gate",   "≥ 95% of ATR",  "Entry penalised −0.40R (soft, not hard mask)"],
        ["Max drawdown per episode", "5R total",   "Episode terminates immediately"],
        ["Max drawdown breach penalty", "−2.00R",  "Hard violation reward applied"],
        ["Daily loss limit breach",     "−1.00R",  "Hard violation reward applied"],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 7. REWARD FUNCTION SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
h1("7. Reward Function Summary")

h2("Core Reward")
body(
    "The primary reward signal on a closed trade is the R-multiple outcome scaled "
    "by the agent's running win rate:"
)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("core_reward  =  pnl_r  ×  win_rate")
run.bold = True
run.font.size = Pt(13)
doc.add_paragraph()
body(
    "Win rate starts at 0.5 (Laplace prior) before any trades and updates after each "
    "closed trade. This means the agent is rewarded not just for profitable trades, "
    "but for being consistently profitable — a single lucky win is discounted."
)

doc.add_paragraph()
h2("Hold-Flat Penalty")
body(
    "A penalty of −0.025R is applied every bar the agent holds flat (no open position). "
    "This is critical for trade frequency: the cumulative daily cost of staying flat "
    "(≈78 bars × 0.025 = 1.95R/day) exceeds the expected cost of a typical loss, "
    "forcing the agent to trade 4–5 times per day rather than staying idle."
)

doc.add_paragraph()
h2("Reward Component Hierarchy")
add_table(
    ["Category", "Components", "Purpose"],
    [
        ["Core outcome",   "pnl_r × win_rate",                          "Primary learning signal"],
        ["Entry shaping",  "Zone/liquidity/rejection bonuses & penalties","Learn WHERE to enter"],
        ["Exit shaping",   "Trail bonuses, gave-back penalties",         "Learn WHEN and HOW to exit"],
        ["Discipline",     "Hold-flat penalty, overtrading penalty",     "Learn HOW OFTEN to trade"],
        ["Violations",     "Drawdown/loss-limit breach penalties",       "Enforce risk guardrails"],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 8. MODEL ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
h1("8. Model Architecture")

add_table(
    ["Component", "Detail"],
    [
        ["Algorithm",        "Recurrent PPO (RecurrentPPO) with LSTM — from sb3_contrib"],
        ["Memory",           "LSTM hidden state persists across bars within an episode"],
        ["Action masking",   "MaskablePPO-style masking applied before sampling"],
        ["Parallelism",      "16 SubprocVecEnv workers during training"],
        ["Normalisation",    "VecNormalize on observations (not rewards) — stats saved per checkpoint"],
        ["Entropy annealing","Entropy coefficient: 0.05 (start) → 0.005 (end) over 1M steps"],
        ["Total timesteps",  "2,000,000 per fold"],
        ["Eval frequency",   "Every 50,000 steps"],
    ]
)

doc.add_paragraph()
h2("Checkpoint Saving — Tiered Phase Gates")
body(
    "Models are saved only when both conditions are met: "
    "(1) composite score ≥ phase threshold, and "
    "(2) ≥ 8 trades evaluated (prevents saving lucky zero-trade runs)."
)
doc.add_paragraph()

add_table(
    ["Phase", "Step Range", "Min Composite Score Required"],
    [
        ["0 — Warmup",      "0 – 400k",    "0.05"],
        ["1 — Post-warmup", "400k – 750k", "0.15"],
        ["2 — Consistency", "750k – 1.1M", "0.25"],
        ["3 — Quality",     "1.1M – 1.5M", "0.35"],
        ["4 — Elite",       "1.5M+",       "0.45"],
    ]
)

doc.add_paragraph()
h2("Composite Score Formula")
body("Composite score (0–1) used for model ranking and save gates:")
doc.add_paragraph()

add_table(
    ["Metric", "Weight", "Reference ('Perfect') Value", "Direction"],
    [
        ["Annualised Sharpe Ratio", "30%", "1.5",  "Higher = better"],
        ["Total PnL (R)",           "25%", "20R",   "Higher = better"],
        ["Win/Loss Ratio",          "25%", "3.0",   "Higher = better"],
        ["Max Drawdown (R)",        "20%", "8R",    "Lower = better"],
    ]
)

doc.add_paragraph()
body(
    "Note: Sharpe ratio is annualised by multiplying the episode-level ratio by √252. "
    "A Sharpe of 1.5 represents an excellent live strategy."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 9. WALK-FORWARD VALIDATION
# ═════════════════════════════════════════════════════════════════════════════
h1("9. Walk-Forward Validation")
body(
    "To test robustness across different market regimes, training uses a rolling "
    "walk-forward framework:"
)
doc.add_paragraph()

add_table(
    ["Parameter", "Value"],
    [
        ["Training window",   "252 trading days (≈ 12 months)"],
        ["Validation window", "26 trading days (≈ 5 weeks)"],
        ["Fold shift",        "26 days — each fold shifts forward by one val window"],
        ["Agent per fold",    "Fresh agent + fresh VecNormalize per fold (no leakage)"],
        ["Output",            "Per-fold Excel journal + Plotly HTML equity curves"],
    ]
)

doc.add_paragraph()
body(
    "After training, test_fold.py loads all saved checkpoints, runs deterministic "
    "evaluation on the held-out test period, and outputs a ranked summary table "
    "with per-checkpoint Plotly candlestick charts showing entries, exits, "
    "SL/TP lines, and the equity curve."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 10. TRADE LIFECYCLE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
h1("10. Complete Trade Lifecycle")

steps = [
    ("Bar-by-bar scan",
     "Each 5-min bar: observation built from ATR state, zone detector, liquidity detector, "
     "trend classifier, order zone engine, and portfolio state (26-40 features)."),
    ("Action masking",
     "Hard safety masks applied — invalid actions removed from agent's choice set."),
    ("Agent decision",
     "LSTM processes observation + hidden state → selects action from masked distribution."),
    ("Entry execution (if ENTER LONG/SHORT)",
     "Stop = zone boundary − ATR buffer.  Target = entry + 75% remaining ATR.  "
     "Contracts sized to risk exactly 1% of capital."),
    ("In-trade management",
     "Each subsequent bar: agent can HOLD (do nothing), EXIT early, or TRAIL STOP.  "
     "Stop/target checked against bar high/low — stop is priority if both hit."),
    ("Trailing stop advancement",
     "When agent selects TRAIL STOP and unrealised ≥ 2R: stop ratchets behind price "
     "by ATR × 25%.  Stop only moves in favour — never reversed."),
    ("Exit triggered",
     "Priority: stop hit → take profit hit → agent EXIT → session end → max drawdown."),
    ("Reward computed",
     "core_reward = pnl_r × win_rate, plus all shaping components.  "
     "Reward stored for PPO update."),
    ("Episode summary",
     "At session end: metrics computed (win rate, PnL, Sharpe, drawdown).  "
     "Composite score checked against phase gate.  Checkpoint saved if gate passes."),
]

for i, (title_text, detail) in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(detail)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ═════════════════════════════════════════════════════════════════════════════
h1("Appendix: Key Configuration Values")

add_table(
    ["Parameter", "Value", "File"],
    [
        ["Bar timeframe",              "5 minutes",          "environment_config.yaml"],
        ["Session type",               "RTH (09:30–16:00 ET)","environment_config.yaml"],
        ["Initial account balance",    "$2,500",              "environment_config.yaml"],
        ["Risk per trade",             "1% of account",       "risk_config.yaml"],
        ["Min / Max contracts",        "1 / 3",               "risk_config.yaml"],
        ["Stop method",                "Zone-based",          "risk_config.yaml"],
        ["Zone buffer (ATR %)",        "3%",                  "risk_config.yaml"],
        ["Take profit method",         "75% remaining ATR",   "risk_config.yaml"],
        ["Minimum R:R",                "1.5:1",               "risk_config.yaml"],
        ["Trail activates at",         "2R unrealised",       "risk_config.yaml"],
        ["Trail aggressive at",        "4R unrealised",       "risk_config.yaml"],
        ["Trail lock-in",              "2R minimum",          "risk_config.yaml"],
        ["Daily loss limit",           "3R",                  "risk_config.yaml"],
        ["Max trades/day",             "10",                  "risk_config.yaml"],
        ["Loss streak threshold",      "3 consecutive",       "risk_config.yaml"],
        ["Hold-flat penalty",          "−0.025R / bar",       "reward_config.yaml"],
        ["Max drawdown (episode)",     "5R",                  "risk_config.yaml"],
        ["ATR exhaustion gate",        "95% of daily ATR",    "risk_config.yaml"],
        ["No-entry window",            "Last 3 bars of RTH",  "risk_config.yaml"],
        ["Min trades to save model",   "8",                   "trading_eval_callback.py"],
    ]
)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(
    os.path.dirname(__file__),
    "RL_Order_Zone_Strategy_Summary.docx"
)
doc.save(out_path)
print(f"Saved: {out_path}")
